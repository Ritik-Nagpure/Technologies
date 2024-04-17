import os
from docx import Document

filedir = str(os.getcwd()) + "\\00_Wordfiles"
os.chdir(filedir)
print(os.getcwd())


def workingdoc():
    # to open a document
    doc = Document('demo.docx')
    # extract paragraphs in a list format
    # print(doc.paragraphs)

    # accessing each paragraph from a list
    print(doc.paragraphs[0].text)
    print(doc.paragraphs[1].text)

    # gets number of types of text formats present in paragraphs groups then into a list format
    print(doc.paragraphs[1].runs)

    # accessing each paragraph text format from a list
    print(doc.paragraphs[1].runs[0].text)
    print(doc.paragraphs[1].runs[1].text)
    print(doc.paragraphs[1].runs[2].text)
    print(doc.paragraphs[1].runs[3].text)
    print(doc.paragraphs[1].runs[4].text)
    print(doc.paragraphs[1].runs[5].text)

    # to check if the format is bold or Italic or underline and returns a true or none

    print(doc.paragraphs[1].runs[0].bold)
    print(doc.paragraphs[1].runs[1].bold)
    print(doc.paragraphs[1].runs[3].italic)
    print(doc.paragraphs[1].runs[5].italic)
    print(doc.paragraphs[1].runs[3].underline)

    # modifying document
    print(doc.paragraphs[0].runs[0].text)
    doc.paragraphs[0].runs[0].text = 'italic and underlined.'
    doc.paragraphs[1].style = 'Title'
    doc.save('demo2.docx')


def creatingdoc():
    doc = Document()
    doc.add_paragraph('Paragraphs').style = 'Title'
    doc.add_paragraph('Hello this is para 1')
    doc.add_paragraph('and this is para 2')
    doc.add_paragraph('finally this is para 3')
    doc.paragraphs[1].add_run('This time we added a new line in same paragraph.').bold = True
    doc.save("created.docx")


def getdocintext(filename):
    doc = Document(filename)
    fulltext = []
    for para in doc.paragraphs:
        fulltext.append(para.text)
    return '\n'.join(fulltext)


print(getdocintext('created.docx'))
